import csv
import io
from pathlib import Path
from typing import Any

from core.exceptions import IngestionError
from core.logging import get_logger
from ingestion.extractors.base import ExtractedDocument

log = get_logger(__name__)


def _pdf_reader(path: Path):
    from pypdf import PdfReader

    return PdfReader(path)


def _convert_pdf_to_images(path: Path, dpi: int = 200):
    from pdf2image import convert_from_path

    return convert_from_path(str(path), dpi=dpi)


def _open_image(path: Path):
    from PIL import Image

    return Image.open(path)


def _ocr_image(image: object) -> str:
    import pytesseract

    if hasattr(image, "close"):
        return pytesseract.image_to_string(image, config="--psm 6").strip()
    with _open_image(image if isinstance(image, Path) else Path(str(image))) as pil_image:
        return pytesseract.image_to_string(pil_image, config="--psm 6").strip()


class PdfExtractor:
    suffixes = (".pdf",)

    def extract(self, path: Path) -> ExtractedDocument:
        reader = _pdf_reader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        page_count = len(pages)
        extracted_text = "\n\n".join(text.strip() for text in pages if text and text.strip())
        if extracted_text.strip():
            return ExtractedDocument(
                text=extracted_text,
                metadata={"page_count": page_count, "extraction_strategy": "pdf-text"},
            )

        ocr_text = _extract_pdf_via_ocr(path, page_count)
        if ocr_text:
            return ExtractedDocument(
                text=ocr_text,
                metadata={
                    "page_count": page_count,
                    "extraction_strategy": "pdf-ocr",
                },
            )

        return ExtractedDocument(text="", metadata={"page_count": page_count, "extraction_strategy": "pdf-empty"})


class ImageExtractor:
    suffixes = (
        ".png",
        ".jpg",
        ".jpeg",
        ".gif",
        ".bmp",
        ".webp",
        ".tiff",
        ".tif",
    )

    def extract(self, path: Path) -> ExtractedDocument:
        text = _ocr_image(path)
        if text:
            return ExtractedDocument(
                text=text,
                metadata={
                    "page_count": 1,
                    "extraction_strategy": "image-ocr",
                },
            )
        return ExtractedDocument(
            text="",
            metadata={"page_count": 1, "extraction_strategy": "image-empty"},
        )


class DocxExtractor:
    suffixes = (".docx",)

    def extract(self, path: Path) -> ExtractedDocument:
        import docx

        document = docx.Document(str(path))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        return ExtractedDocument(
            text="\n\n".join(paragraphs), metadata={"paragraph_count": len(paragraphs)}
        )


class MarkdownExtractor:
    suffixes = (".md",)

    def extract(self, path: Path) -> ExtractedDocument:
        text = path.read_text(encoding="utf-8", errors="replace")
        headings = [
            line.lstrip("# ").strip()
            for line in text.splitlines()
            if line.startswith("#")
        ]
        return ExtractedDocument(text=text, metadata={"headings": headings[:50], "format": "markdown"})


class PlainTextExtractor:
    suffixes = (".txt",)

    def extract(self, path: Path) -> ExtractedDocument:
        return ExtractedDocument(text=path.read_text(encoding="utf-8", errors="replace"))


class CodeExtractor:
    suffixes = (
        ".c",
        ".cc",
        ".cpp",
        ".cs",
        ".go",
        ".graphql",
        ".h",
        ".hpp",
        ".java",
        ".js",
        ".jsx",
        ".kt",
        ".php",
        ".proto",
        ".py",
        ".rb",
        ".rs",
        ".sh",
        ".sql",
        ".swift",
        ".toml",
        ".ts",
        ".tsx",
        ".xml",
        ".yaml",
        ".yml",
    )

    def extract(self, path: Path) -> ExtractedDocument:
        return ExtractedDocument(
            text=path.read_text(encoding="utf-8", errors="strict"),
            metadata={"format": "code", "language": path.suffix.lower().lstrip(".")},
        )


class CsvExtractor:
    suffixes = (".csv",)

    def extract(self, path: Path) -> ExtractedDocument:
        raw = path.read_text(encoding="utf-8", errors="replace")
        rows = list(csv.reader(io.StringIO(raw)))
        if not rows:
            return ExtractedDocument(text="", metadata={"row_count": 0})
        header, body = rows[0], rows[1:]
        # row-per-line "column: value" rendering keeps cell/column association
        # visible to both FTS and the LLM
        lines = [
            "; ".join(f"{h}: {v}" for h, v in zip(header, row, strict=False))
            for row in body
        ]
        return ExtractedDocument(
            text="\n".join(lines), metadata={"row_count": len(body), "columns": header}
        )


class HtmlExtractor:
    suffixes = (".html", ".htm")

    def extract(self, path: Path) -> ExtractedDocument:
        from bs4 import BeautifulSoup
        from bs4.element import Tag

        soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        title = soup.title.string.strip() if soup.title and soup.title.string else None
        metadata: dict[str, str | None] = {"html_title": title}
        for meta_name in ("source-url", "confluence-page-id", "confluence-space-key"):
            tag = soup.find("meta", attrs={"name": meta_name})
            if tag is not None:
                content = tag.get("content")
                if isinstance(content, str) and content.strip():
                    metadata[meta_name] = content.strip()
        lines: list[str] = []
        headings: list[str] = []
        for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "table"]):
            if not isinstance(element, Tag) or _has_parent(element, {"li", "pre", "table"}):
                continue
            name = element.name.lower()
            text = element.get_text(separator=" ", strip=True)
            if not text:
                continue
            if name.startswith("h") and len(name) == 2 and name[1].isdigit():
                level = min(max(int(name[1]), 1), 6)
                headings.append(text)
                lines.append(f"{'#' * level} {text}")
            elif name == "li":
                lines.append(f"- {text}")
            elif name == "pre":
                lines.append(f"```\n{text}\n```")
            elif name == "table":
                table_lines = _table_lines(element)
                if table_lines:
                    lines.extend(table_lines)
            else:
                lines.append(text)
        metadata["headings"] = headings[:50]
        metadata["format"] = "html"
        return ExtractedDocument(text="\n\n".join(lines).strip(), metadata=metadata)


_EXTRACTORS = [
    PdfExtractor(),
    ImageExtractor(),
    DocxExtractor(),
    MarkdownExtractor(),
    PlainTextExtractor(),
    CodeExtractor(),
    CsvExtractor(),
    HtmlExtractor(),
]

SUPPORTED_SUFFIXES: frozenset[str] = frozenset(
    s for e in _EXTRACTORS for s in e.suffixes
)


def extract_text(path: Path) -> ExtractedDocument:
    suffix = path.suffix.lower()
    for extractor in _EXTRACTORS:
        if suffix in extractor.suffixes:
            doc = extractor.extract(path)
            if not doc.text.strip():
                raise IngestionError(f"No extractable text in {path.name}")
            return doc
    raise IngestionError(f"Unsupported file type: {suffix}")


def _extract_pdf_via_ocr(path: Path, page_count: int) -> str:
    try:
        images = _convert_pdf_to_images(path)
    except Exception:
        return ""

    chunks: list[str] = []
    for index, image in enumerate(images, start=1):
        try:
            text = _ocr_image(image)
        except Exception as exc:
            log.warning("pdf_ocr_page_failed", page=index, error=str(exc)[:500])
            continue
        if text:
            chunks.append(f"Page {index}: {text.strip()}")
    if not chunks and page_count > 0:
        return ""
    return "\n\n".join(chunks)


def _has_parent(element: Any, names: set[str]) -> bool:
    parent = getattr(element, "parent", None)
    while parent is not None:
        name = getattr(parent, "name", None)
        if isinstance(name, str) and name.lower() in names:
            return True
        parent = getattr(parent, "parent", None)
    return False


def _table_lines(table: Any) -> list[str]:
    rows: list[str] = []
    for tr in table.find_all("tr"):
        cells = [cell.get_text(" ", strip=True) for cell in tr.find_all(["th", "td"])]
        cells = [cell for cell in cells if cell]
        if cells:
            rows.append(" | ".join(cells))
    return rows
