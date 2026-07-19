from __future__ import annotations

import asyncio
import hashlib
import io
import re
import uuid
import zipfile
from collections.abc import AsyncIterator
from dataclasses import dataclass, replace
from typing import Self, cast
from urllib.parse import urlparse

import httpx
from defusedxml import ElementTree
from sqlalchemy.ext.asyncio import AsyncSession

from core.config import Settings
from core.exceptions import AuthorizationError, NotFoundError, ProviderError, ValidationError
from embeddings.base import EmbeddingProvider
from ingestion.chunkers.recursive import RecursiveChunker
from ingestion.queue import IngestionQueue
from knowledgebase.source_metadata import normalize_source_metadata
from models import Document, DocumentStatus, KnowledgeBase, User
from repositories.audit import AuditLogRepository
from repositories.knowledge import DocumentRepository, KnowledgeBaseRepository
from retrieval.context import RetrievedChunk
from services.document_service import DocumentService

JsonDict = dict[str, object]
CHUNK_STRATEGY_VERSION = "jira-relationship-comments-attachments-v5"
JIRA_FIELDS = (
    "summary,status,issuetype,priority,labels,components,assignee,reporter,"
    "updated,created,description,project,parent,subtasks,issuelinks,attachment,comment"
)


@dataclass(frozen=True, slots=True)
class JiraBoard:
    id: int
    name: str
    type: str
    url: str


@dataclass(frozen=True, slots=True)
class JiraIssueRef:
    key: str
    title: str
    relationship: str
    status: str | None = None


@dataclass(frozen=True, slots=True)
class JiraComment:
    id: str
    author: str
    body: str
    created_at: str | None
    updated_at: str | None


@dataclass(frozen=True, slots=True)
class JiraAttachment:
    id: str
    filename: str
    mime_type: str | None
    size_bytes: int | None
    content_url: str | None
    created_at: str | None
    author: str | None
    extracted_text: str = ""


@dataclass(frozen=True, slots=True)
class JiraIssue:
    id: str
    key: str
    title: str
    url: str
    issue_type: str | None
    status: str | None
    status_category: str | None
    status_category_key: str | None
    priority: str | None
    labels: list[str]
    components: list[str]
    assignee: str | None
    assignee_email: str | None
    assignee_account_id: str | None
    reporter: str | None
    reporter_email: str | None
    reporter_account_id: str | None
    created_at: str | None
    updated_at: str | None
    description: str
    project_key: str | None
    project_name: str | None
    parent: JiraIssueRef | None = None
    related_issues: tuple[JiraIssueRef, ...] = ()
    child_issues: tuple[JiraIssueRef, ...] = ()
    comments: tuple[JiraComment, ...] = ()
    comment_total: int | None = None
    attachments: tuple[JiraAttachment, ...] = ()


@dataclass(frozen=True, slots=True)
class SyncedJiraDocument:
    issue_id: str
    issue_key: str
    title: str
    url: str
    status: str | None
    updated_at: str | None
    document_id: uuid.UUID
    document_status: str
    action: str


@dataclass(frozen=True, slots=True)
class JiraSyncResult:
    knowledge_base_id: uuid.UUID
    knowledge_base_name: str
    project_key: str
    board_id: int
    board_name: str
    total_issues: int
    created: int
    updated: int
    skipped: int
    documents: list[SyncedJiraDocument]


class JiraClient:
    def __init__(self, settings: Settings) -> None:
        self._settings = settings
        self._base_url = normalize_jira_base_url(settings.jira_base_url)
        self._client: httpx.AsyncClient | None = None

    async def __aenter__(self) -> Self:
        token = jira_token(self._settings)
        if not token:
            raise ValidationError("Jira API token is not configured")

        mode = self._settings.jira_auth_mode.strip().lower()
        if mode not in {"auto", "basic", "bearer"}:
            raise ValidationError("JIRA_AUTH_MODE must be auto, basic, or bearer")

        auth: httpx.Auth | None = None
        headers = {"Accept": "application/json", "User-Agent": "cvum-knowledge-hub/0.1"}
        email = jira_email(self._settings)
        if mode == "basic" or (mode == "auto" and email):
            if not email:
                raise ValidationError("JIRA_EMAIL or CONFLUENCE_EMAIL is required for Jira basic auth")
            auth = httpx.BasicAuth(email, token)
        else:
            headers["Authorization"] = f"Bearer {token}"

        self._client = httpx.AsyncClient(
            base_url=f"{self._base_url}/",
            auth=auth,
            headers=headers,
            timeout=self._settings.jira_request_timeout_seconds,
            follow_redirects=True,
        )
        return self

    async def __aexit__(self, *_exc: object) -> None:
        if self._client is not None:
            await self._client.aclose()

    async def get_board(self, board_id: int) -> JiraBoard:
        payload = await self._get(f"rest/agile/1.0/board/{board_id}")
        board_id_value = _int(payload.get("id"))
        if board_id_value is None:
            raise ProviderError("Jira board response is missing id")
        return JiraBoard(
            id=board_id_value,
            name=_str(payload.get("name")) or f"Board {board_id_value}",
            type=_str(payload.get("type")) or "unknown",
            url=f"{self._base_url}/jira/software/c/projects/{self._settings.jira_project_key}/boards/{board_id_value}",
        )

    async def get_issue(self, issue_key: str) -> JiraIssue:
        payload = await self._get(
            f"rest/api/3/issue/{issue_key}",
            params={"fields": JIRA_FIELDS},
        )
        return self._parse_issue(payload)

    async def get_child_issues(self, issue_key: str, *, max_issues: int = 50) -> list[JiraIssue]:
        params: dict[str, str | int] = {
            "jql": f'parent = "{issue_key}" ORDER BY updated DESC',
            "maxResults": min(max_issues, 100),
            "fields": JIRA_FIELDS,
        }
        payload = await self._get("rest/api/3/search/jql", params=params)
        return [self._parse_issue(row) for row in _list_of_dicts(payload.get("issues"))[:max_issues]]

    async def iter_board_issues(self, board_id: int, max_issues: int | None = None) -> AsyncIterator[JiraIssue]:
        limit = max(1, min(self._settings.jira_issue_limit, 100))
        if max_issues is not None:
            limit = min(limit, max(1, max_issues))
        path = f"rest/software/1.0/board/{board_id}/issue"
        params: dict[str, str | int] = {
            "maxResults": limit,
            "fields": JIRA_FIELDS,
        }
        if self._settings.jira_project_key.strip():
            params["jql"] = f"project = {self._settings.jira_project_key.strip()}"
        fetched = 0

        while max_issues is None or fetched < max_issues:
            payload = await self._get(path, params=params)
            issues = _list_of_dicts(payload.get("issues"))
            for row in issues:
                yield self._parse_issue(row)
                fetched += 1
                if max_issues is not None and fetched >= max_issues:
                    return

            token = _str(payload.get("nextPageToken"))
            if token:
                params = {
                    "maxResults": limit,
                    "nextPageToken": token,
                    "fields": JIRA_FIELDS,
                }
                if self._settings.jira_project_key.strip():
                    params["jql"] = f"project = {self._settings.jira_project_key.strip()}"
                continue

            start = _int(payload.get("startAt"))
            total = _int(payload.get("total"))
            max_results = _int(payload.get("maxResults")) or limit
            if start is not None and total is not None and start + max_results < total:
                params = {
                    "maxResults": limit,
                    "startAt": start + max_results,
                    "fields": JIRA_FIELDS,
                }
                if self._settings.jira_project_key.strip():
                    params["jql"] = f"project = {self._settings.jira_project_key.strip()}"
                continue
            return

    async def iter_project_issues(self, project_key: str, max_issues: int | None = None) -> AsyncIterator[JiraIssue]:
        key = project_key.strip()
        if not key:
            raise ValidationError("JIRA_PROJECT_KEY is required for project issue sync")
        limit = max(1, min(self._settings.jira_issue_limit, 100))
        if max_issues is not None:
            limit = min(limit, max(1, max_issues))
        params: dict[str, str | int] = {
            "jql": f"project = {key} ORDER BY updated DESC",
            "maxResults": limit,
            "fields": JIRA_FIELDS,
        }
        fetched = 0
        start_at = 0
        use_legacy_search = False

        while max_issues is None or fetched < max_issues:
            if use_legacy_search:
                payload = await self._get("rest/api/3/search", params={**params, "startAt": start_at})
            else:
                try:
                    payload = await self._get("rest/api/3/search/jql", params=params)
                except NotFoundError:
                    use_legacy_search = True
                    payload = await self._get("rest/api/3/search", params={**params, "startAt": start_at})

            issues = _list_of_dicts(payload.get("issues"))
            for row in issues:
                yield self._parse_issue(row)
                fetched += 1
                if max_issues is not None and fetched >= max_issues:
                    return

            token = _str(payload.get("nextPageToken"))
            if token and not use_legacy_search:
                params = {
                    "jql": f"project = {key} ORDER BY updated DESC",
                    "maxResults": limit,
                    "nextPageToken": token,
                    "fields": JIRA_FIELDS,
                }
                continue

            start = _int(payload.get("startAt"))
            total = _int(payload.get("total"))
            max_results = _int(payload.get("maxResults")) or limit
            if start is not None and total is not None and start + max_results < total:
                start_at = start + max_results
                continue
            return

    async def hydrate_issue(self, issue: JiraIssue) -> JiraIssue:
        comments = issue.comments
        if self._settings.jira_include_comments:
            needs_comment_page = issue.comment_total is None or issue.comment_total > len(comments)
            if needs_comment_page:
                try:
                    comments = tuple(
                        await self.get_comments(
                            issue.key,
                            max_comments=self._settings.jira_max_comments_per_issue,
                        )
                    )
                except (AuthorizationError, NotFoundError, ProviderError):
                    comments = issue.comments

        attachments: list[JiraAttachment] = []
        for attachment in issue.attachments[: self._settings.jira_max_attachments_per_issue]:
            if not self._settings.jira_extract_attachments or not _attachment_is_extractable(attachment):
                attachments.append(attachment)
                continue
            if attachment.size_bytes is not None and attachment.size_bytes > self._settings.jira_attachment_max_bytes:
                attachments.append(attachment)
                continue
            try:
                payload = await self._get_bytes(attachment.content_url or "")
                extracted = _extract_attachment_text(attachment.filename, payload)
            except (
                AuthorizationError,
                NotFoundError,
                ProviderError,
                ImportError,
                OSError,
                RuntimeError,
                ValueError,
                zipfile.BadZipFile,
            ):
                extracted = ""
            attachments.append(replace(attachment, extracted_text=extracted))
        return replace(issue, comments=comments, attachments=tuple(attachments))

    async def get_comments(self, issue_key: str, *, max_comments: int) -> list[JiraComment]:
        comments: list[JiraComment] = []
        start_at = 0
        page_size = min(max(max_comments, 1), 100)
        while len(comments) < max_comments:
            payload = await self._get(
                f"rest/api/3/issue/{issue_key}/comment",
                params={"startAt": start_at, "maxResults": page_size, "orderBy": "created"},
            )
            rows = _list_of_dicts(payload.get("comments"))
            for row in rows:
                comment = _parse_comment(row, fallback_id=str(len(comments) + 1))
                if comment is None:
                    continue
                comments.append(comment)
                if len(comments) >= max_comments:
                    return comments
            total = _int(payload.get("total")) or 0
            start_at += len(rows)
            if not rows or start_at >= total:
                return comments
        return comments

    async def _get_bytes(self, path_or_url: str) -> bytes:
        if self._client is None:
            raise RuntimeError("JiraClient must be used as an async context manager")
        if not path_or_url:
            raise ValueError("Attachment URL is empty")
        try:
            response = await self._client.get(path_or_url)
        except httpx.HTTPError as exc:
            raise ProviderError("Jira attachment request failed") from exc
        if response.status_code == 403:
            raise AuthorizationError("Jira attachment is not visible to this account")
        if response.status_code == 404:
            raise NotFoundError("Jira attachment was not found")
        if response.status_code >= 400:
            raise ProviderError(f"Jira attachment request failed with HTTP {response.status_code}")
        return response.content

    async def _get(self, path: str, params: dict[str, str | int] | None = None) -> JsonDict:
        if self._client is None:
            raise RuntimeError("JiraClient must be used as an async context manager")
        try:
            response = await self._client.get(path, params=params)
        except httpx.HTTPError as exc:
            raise ProviderError("Jira read request failed") from exc

        if response.status_code == 401:
            raise ProviderError("Jira authentication failed")
        if response.status_code == 403:
            raise AuthorizationError("Jira token does not have permission to view this board")
        if response.status_code == 404:
            raise NotFoundError("Jira board or issue resource was not found")
        if response.status_code >= 400:
            raise ProviderError(f"Jira read request failed with HTTP {response.status_code}")

        payload = response.json()
        if not isinstance(payload, dict):
            raise ProviderError("Jira returned an unexpected response shape")
        return cast(JsonDict, payload)

    def _parse_issue(self, row: JsonDict) -> JiraIssue:
        fields = _dict(row.get("fields"))
        project = _dict(fields.get("project"))
        status = _dict(fields.get("status"))
        status_category = _dict(status.get("statusCategory"))
        issue_type = _dict(fields.get("issuetype"))
        priority = _dict(fields.get("priority"))
        components = _list_of_dicts(fields.get("components"))
        assignee = _dict(fields.get("assignee"))
        reporter = _dict(fields.get("reporter"))
        parent = _parse_issue_ref(fields.get("parent"), relationship="parent")
        related_issues = _parse_issue_links(fields.get("issuelinks"))
        subtasks = tuple(
            ref
            for item in _list(fields.get("subtasks"))
            if (ref := _parse_issue_ref(item, relationship="child")) is not None
        )
        attachments = tuple(
            attachment
            for item in _list(fields.get("attachment"))
            if (attachment := _parse_attachment(item)) is not None
        )
        comment_page = _dict(fields.get("comment"))
        comments = tuple(
            comment
            for index, item in enumerate(_list_of_dicts(comment_page.get("comments")), start=1)
            if (comment := _parse_comment(item, fallback_id=str(index))) is not None
        )
        comment_total = _int(comment_page.get("total")) if comment_page else None
        key = _required_str(row, "key")
        issue_id = _required_str(row, "id")
        return JiraIssue(
            id=issue_id,
            key=key,
            title=_str(fields.get("summary")) or key,
            url=f"{self._base_url}/browse/{key}",
            issue_type=_str(issue_type.get("name")),
            status=_str(status.get("name")),
            status_category=_str(status_category.get("name")),
            status_category_key=_str(status_category.get("key")),
            priority=_str(priority.get("name")),
            labels=[str(label) for label in _list(fields.get("labels")) if isinstance(label, str)],
            components=[name for item in components if (name := _str(item.get("name")))],
            assignee=_str(assignee.get("displayName")),
            assignee_email=_str(assignee.get("emailAddress")),
            assignee_account_id=_str(assignee.get("accountId")),
            reporter=_str(reporter.get("displayName")),
            reporter_email=_str(reporter.get("emailAddress")),
            reporter_account_id=_str(reporter.get("accountId")),
            created_at=_str(fields.get("created")),
            updated_at=_str(fields.get("updated")),
            description=adf_to_text(fields.get("description")),
            project_key=_str(project.get("key")),
            project_name=_str(project.get("name")),
            parent=parent,
            related_issues=related_issues,
            child_issues=subtasks,
            comments=comments,
            comment_total=comment_total,
            attachments=attachments,
        )


class JiraSyncService:
    def __init__(
        self,
        *,
        db: AsyncSession,
        settings: Settings,
        embedder: EmbeddingProvider,
        queue: IngestionQueue,
    ) -> None:
        self._db = db
        self._settings = settings
        self._docs = DocumentRepository(db)
        self._kbs = KnowledgeBaseRepository(db)
        self._audit = AuditLogRepository(db)
        self._document_service = DocumentService(
            db=db, settings=settings, embedder=embedder, queue=queue
        )

    async def sync_board(
        self,
        *,
        user: User,
        kb_id: uuid.UUID | None = None,
        max_issues: int | None = None,
    ) -> JiraSyncResult:
        kb: KnowledgeBase | None = None
        try:
            self._validate_config()
            kb = await self._resolve_kb(user=user, kb_id=kb_id)
            limit = _sync_limit(max_issues, self._settings.jira_sync_max_issues)
            documents: list[SyncedJiraDocument] = []

            async with JiraClient(self._settings) as client:
                if self._settings.jira_board_id:
                    board = await client.get_board(self._settings.jira_board_id)
                    issues = client.iter_board_issues(board.id, max_issues=limit)
                else:
                    board = JiraBoard(
                        id=0,
                        name=f"{self._settings.jira_project_key.strip()} Project",
                        type="project",
                        url=f"{normalize_jira_base_url(self._settings.jira_base_url)}/jira/core/projects/{self._settings.jira_project_key.strip()}/board",
                    )
                    issues = client.iter_project_issues(self._settings.jira_project_key, max_issues=limit)

                issue_rows = [
                    issue async for issue in issues if _jira_issue_allowed(issue, self._settings)
                ]
                issue_rows = _attach_reverse_children(issue_rows)
                issue_rows = await _hydrate_issues(
                    client,
                    issue_rows,
                    concurrency=self._settings.jira_hydration_concurrency,
                )
                for issue in issue_rows:
                    documents.append(await self._sync_issue(user=user, kb=kb, board=board, issue=issue))
        except Exception as exc:
            await self._db.rollback()
            self._audit.record(
                action="jira.sync",
                resource_type="knowledge_base",
                resource_id=str(kb.id) if kb else None,
                org_id=user.organization_id,
                actor_id=user.id,
                detail=f"0 created, 0 updated, 0 skipped, 1 failed: {type(exc).__name__}: {str(exc)[:180]}",
            )
            await self._db.commit()
            raise

        created = sum(1 for doc in documents if doc.action == "created")
        updated = sum(1 for doc in documents if doc.action == "updated")
        skipped = sum(1 for doc in documents if doc.action == "skipped")
        self._audit.record(
            action="jira.sync",
            resource_type="knowledge_base",
            resource_id=str(kb.id),
            org_id=user.organization_id,
            actor_id=user.id,
            detail=f"{created} created, {updated} updated, {skipped} skipped from {board.name}",
        )
        await self._db.commit()
        return JiraSyncResult(
            knowledge_base_id=kb.id,
            knowledge_base_name=kb.name,
            project_key=self._settings.jira_project_key,
            board_id=board.id,
            board_name=board.name,
            total_issues=len(documents),
            created=created,
            updated=updated,
            skipped=skipped,
            documents=documents,
        )

    async def _sync_issue(
        self,
        *,
        user: User,
        kb: KnowledgeBase,
        board: JiraBoard,
        issue: JiraIssue,
    ) -> SyncedJiraDocument:
        existing = await self._docs.get_by_metadata_value(kb.id, "jira_issue_id", issue.id)
        metadata = _issue_metadata(board=board, issue=issue)
        if existing is not None and _is_current(existing, metadata):
            if (existing.doc_metadata or {}) != metadata:
                await self._docs.update_metadata(existing.id, metadata)
            await self._docs.soft_delete_metadata_duplicates(
                kb.id, "jira_issue_id", issue.id, existing.id
            )
            return SyncedJiraDocument(
                issue_id=issue.id,
                issue_key=issue.key,
                title=issue.title,
                url=issue.url,
                status=issue.status,
                updated_at=issue.updated_at,
                document_id=existing.id,
                document_status=existing.status.value,
                action="skipped",
            )

        content = _render_issue_markdown(board=board, issue=issue).encode("utf-8")
        doc = await self._document_service.create_from_bytes(
            user=user,
            kb_id=kb.id,
            filename=_issue_filename(issue),
            content=content,
            existing_document_id=existing.id if existing else None,
            title=f"{issue.key}: {issue.title}",
            metadata=metadata,
            audit_action="jira.issue.sync",
        )
        await self._docs.soft_delete_metadata_duplicates(kb.id, "jira_issue_id", issue.id, doc.id)
        return SyncedJiraDocument(
            issue_id=issue.id,
            issue_key=issue.key,
            title=issue.title,
            url=issue.url,
            status=issue.status,
            updated_at=issue.updated_at,
            document_id=doc.id,
            document_status=doc.status.value,
            action="updated" if existing else "created",
        )

    async def _resolve_kb(self, *, user: User, kb_id: uuid.UUID | None) -> KnowledgeBase:
        if kb_id is not None:
            kb = await self._kbs.get(kb_id, user.organization_id)
            if kb is None:
                raise NotFoundError("Knowledge base not found")
            return kb

        existing = [
            kb
            for kb in await self._kbs.list_by_org(user.organization_id)
            if kb.name == self._settings.jira_default_kb_name
        ]
        if existing:
            return existing[0]

        kb = KnowledgeBase(
            organization_id=user.organization_id,
            name=self._settings.jira_default_kb_name,
            description=f"Read-only Jira sync for board {self._settings.jira_board_id}.",
            embedding_model=self._settings.embedding_model,
            embedding_dimensions=self._settings.embedding_dimensions,
        )
        self._kbs.add(kb)
        await self._db.commit()
        return kb

    def _validate_config(self) -> None:
        if not self._settings.jira_base_url.strip():
            raise ValidationError("JIRA_BASE_URL is not configured")
        if not self._settings.jira_board_id and not self._settings.jira_project_key.strip():
            raise ValidationError("JIRA_BOARD_ID or JIRA_PROJECT_KEY is not configured")
        if not jira_token(self._settings):
            raise ValidationError("JIRA_API_TOKEN or CONFLUENCE_API_TOKEN is not configured")
        if self._settings.jira_auth_mode.strip().lower() == "basic" and not jira_email(self._settings):
            raise ValidationError("JIRA_EMAIL or CONFLUENCE_EMAIL is required when JIRA_AUTH_MODE=basic")


def jira_config_status(settings: Settings) -> JsonDict:
    base_url = settings.jira_base_url.strip()
    token = jira_token(settings)
    email = jira_email(settings)
    mode = settings.jira_auth_mode.strip().lower()
    using_fallback = bool(token and not settings.jira_api_token.strip())
    return {
        "configured": bool(
            base_url
            and (settings.jira_board_id or settings.jira_project_key.strip())
            and token
            and (email or mode != "basic")
        ),
        "read_only": True,
        "base_url": normalize_jira_base_url(base_url) if base_url else None,
        "project_key": settings.jira_project_key,
        "board_id": settings.jira_board_id,
        "default_kb_name": settings.jira_default_kb_name,
        "auth_mode": mode,
        "email_configured": bool(email),
        "token_configured": bool(token),
        "using_atlassian_fallback_credentials": using_fallback,
        "requires_email": mode == "basic" or (mode == "auto" and bool(token) and not email),
    }


async def retrieve_live_jira_relationships(
    *,
    settings: Settings,
    issue_key: str,
    query: str,
    limit: int,
) -> list[RetrievedChunk]:
    """Read an exact Jira issue family so epic follow-ups include current comments and files."""

    async with JiraClient(settings) as client:
        root = await client.get_issue(issue_key)
        children = await client.get_child_issues(issue_key, max_issues=50)
        related = _attach_reverse_children([root, *children])
        hydrated = await _hydrate_issues(
            client,
            related,
            concurrency=settings.jira_hydration_concurrency,
        )

    board = JiraBoard(
        id=settings.jira_board_id,
        name=f"{settings.jira_project_key} live Jira",
        type="live",
        url=f"{normalize_jira_base_url(settings.jira_base_url)}/browse/{issue_key}",
    )
    chunker = RecursiveChunker()
    query_terms = _evidence_terms(query)
    candidates: list[RetrievedChunk] = []
    for issue in hydrated:
        rendered = _render_issue_markdown(board=board, issue=issue)
        metadata = _issue_metadata(board=board, issue=issue)
        metadata["retrieval_origin"] = "live_jira_relationship"
        metadata["chunk_profile"] = CHUNK_STRATEGY_VERSION
        for text_chunk in chunker.chunk(
            rendered,
            chunk_size=settings.jira_chunk_size_tokens,
            overlap=settings.jira_chunk_overlap_tokens,
        ):
            content = _live_chunk_content(issue=issue, content=text_chunk.content)
            chunk_terms = _evidence_terms(content)
            overlap = len(query_terms & chunk_terms) / max(len(query_terms), 1)
            is_root = issue.key.upper() == issue_key.upper()
            score = min(
                1.0,
                (0.7 if is_root else 0.62)
                + (0.38 * overlap)
                + _evidence_intent_boost(query_terms=query_terms, chunk_terms=chunk_terms),
            )
            document_id = uuid.uuid5(uuid.NAMESPACE_URL, issue.url)
            candidates.append(
                RetrievedChunk(
                    chunk_id=uuid.uuid5(
                        uuid.NAMESPACE_URL,
                        f"{issue.url}#live-jira-{text_chunk.ordinal}",
                    ),
                    document_id=document_id,
                    document_title=f"{issue.key}: {issue.title}",
                    content=content,
                    metadata={**metadata, "chunk_ordinal": text_chunk.ordinal},
                    sparse_score=score,
                    score=score,
                )
            )

    ranked = sorted(candidates, key=lambda item: item.score, reverse=True)
    selected: list[RetrievedChunk] = []
    per_document: dict[uuid.UUID, int] = {}
    for chunk in ranked:
        cap = 4 if str(chunk.metadata.get("jira_issue_key", "")).upper() == issue_key.upper() else 2
        if per_document.get(chunk.document_id, 0) >= cap:
            continue
        selected.append(chunk)
        per_document[chunk.document_id] = per_document.get(chunk.document_id, 0) + 1
        if len(selected) >= limit:
            break
    return selected


def _live_chunk_content(*, issue: JiraIssue, content: str) -> str:
    prefix = ["Source type: jira", f"Title: {issue.key}: {issue.title}"]
    if issue.parent is not None:
        prefix.append(f"Parent Jira issue: {issue.parent.key}")
    if issue.child_issues:
        prefix.append(
            "Connected Jira issues: " + ", ".join(item.key for item in issue.child_issues[:50])
        )
    return "\n".join([*prefix, "", content.strip()])


def _evidence_terms(value: str) -> set[str]:
    stopwords = {
        "about",
        "and",
        "from",
        "give",
        "have",
        "into",
        "jira",
        "that",
        "the",
        "this",
        "what",
        "where",
        "which",
        "with",
    }
    return {
        token.lower()
        for token in re.findall(r"[A-Za-z][A-Za-z0-9_.-]{2,}", value)
        if token.lower() not in stopwords
    }


def _evidence_intent_boost(*, query_terms: set[str], chunk_terms: set[str]) -> float:
    boost = 0.0
    if query_terms & {"cpu", "ram", "memory", "server", "instance", "shape", "specification"} and chunk_terms & {
        "cpu",
        "ocpu",
        "vcpu",
        "ram",
        "memory",
        "server",
        "instance",
        "shape",
        "compute",
    }:
        boost += 0.16
    if query_terms & {"benchmark", "benchmarking", "load", "performance", "throughput", "latency"} and chunk_terms & {
        "benchmark",
        "load",
        "performance",
        "throughput",
        "latency",
        "meters",
        "aws",
        "oci",
    }:
        boost += 0.12
    return boost


def _sync_limit(requested: int | None, configured: int) -> int | None:
    if requested is not None:
        return requested
    return configured if configured > 0 else None


def normalize_jira_base_url(raw_url: str) -> str:
    raw = raw_url.strip().rstrip("/")
    if not raw:
        raise ValidationError("JIRA_BASE_URL is not configured")
    parsed = urlparse(raw)
    if not parsed.scheme or not parsed.netloc:
        raise ValidationError("JIRA_BASE_URL must be an absolute URL")
    return f"{parsed.scheme}://{parsed.netloc}"


def jira_token(settings: Settings) -> str:
    return settings.jira_api_token.strip() or settings.confluence_api_token.strip()


def jira_email(settings: Settings) -> str:
    return settings.jira_email.strip() or settings.confluence_email.strip()


def adf_to_text(value: object) -> str:
    """Render Jira ADF as retrieval-friendly Markdown without account-id noise."""

    return _clean_jira_description(_render_adf(value)).strip()


def _render_adf(value: object, *, list_depth: int = 0, ordered: bool = False) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        return "".join(_render_adf(item, list_depth=list_depth, ordered=ordered) for item in value)
    if not isinstance(value, dict):
        return ""

    node_type = str(value.get("type") or "")
    attrs = _dict(value.get("attrs"))
    children = _list(value.get("content"))
    if node_type == "text":
        return _str(value.get("text")) or ""
    if node_type == "hardBreak":
        return "\n"
    if node_type == "heading":
        level = min(max(_int(attrs.get("level")) or 2, 1), 6)
        return f"{'#' * level} {_inline_adf(children).strip()}\n\n"
    if node_type == "paragraph":
        text = _inline_adf(children).strip()
        return f"{text}\n\n" if text else ""
    if node_type in {"bulletList", "orderedList"}:
        is_ordered = node_type == "orderedList"
        return "".join(
            _render_adf(child, list_depth=list_depth, ordered=is_ordered) for child in children
        ) + "\n"
    if node_type == "listItem":
        rendered = "".join(
            _render_adf(child, list_depth=list_depth + 1, ordered=ordered) for child in children
        ).strip()
        prefix = "1." if ordered else "-"
        indent = "  " * list_depth
        lines = rendered.splitlines() or [""]
        continuation = "\n".join(f"{indent}  {line}" for line in lines[1:] if line.strip())
        return f"{indent}{prefix} {lines[0]}\n{continuation}\n"
    if node_type == "codeBlock":
        language = str(attrs.get("language") or "").strip()
        return f"```{language}\n{_inline_adf(children).strip()}\n```\n\n"
    if node_type == "blockquote":
        body = _render_adf(children).strip()
        return "\n".join(f"> {line}" for line in body.splitlines()) + "\n\n"
    if node_type == "table":
        return _render_adf_table(children)
    if node_type in {"panel", "expand", "nestedExpand", "doc"}:
        return _render_adf(children, list_depth=list_depth, ordered=ordered)
    if node_type == "status":
        return str(attrs.get("text") or "").strip()
    if node_type == "mention":
        label = str(attrs.get("text") or attrs.get("displayName") or "user").strip().lstrip("@")
        return f"@{label}"
    if node_type == "inlineCard":
        return str(attrs.get("title") or "linked content").strip()
    if node_type in {"media", "mediaGroup", "mediaSingle", "rule"}:
        return "\n"
    return _render_adf(children, list_depth=list_depth, ordered=ordered)


def _inline_adf(children: list[object]) -> str:
    return "".join(_render_adf(child) for child in children).replace("\n\n", "\n")


def _render_adf_table(children: list[object]) -> str:
    rows: list[list[str]] = []
    for row in children:
        row_dict = _dict(row)
        if row_dict.get("type") != "tableRow":
            continue
        cells = [
            _render_adf(_list(_dict(cell).get("content"))).strip().replace("\n", " ")
            for cell in _list(row_dict.get("content"))
        ]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    lines = [" | ".join(normalized[0]), " | ".join(["---"] * width)]
    lines.extend(" | ".join(row) for row in normalized[1:])
    return "\n".join(lines) + "\n\n"


def _clean_jira_description(value: str) -> str:
    value = value.replace("\u200b", "").replace("\ufeff", "")
    lines: list[str] = []
    previous = ""
    for raw_line in value.splitlines():
        line = re.sub(r"[ \t]+", " ", raw_line).strip()
        lowered = line.lower()
        if not line:
            if lines and lines[-1]:
                lines.append("")
            continue
        if lowered in {"no description", "no description provided", "n/a", "none"}:
            continue
        if any(
            marker in lowered
            for marker in (
                "this issue was automatically generated",
                "do not edit this issue manually",
                "sent from jira",
                "generated by jira automation",
            )
        ):
            continue
        if line == previous:
            continue
        lines.append(line)
        previous = line
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


async def _hydrate_issues(
    client: JiraClient,
    issues: list[JiraIssue],
    *,
    concurrency: int,
) -> list[JiraIssue]:
    semaphore = asyncio.Semaphore(max(1, concurrency))

    async def hydrate(issue: JiraIssue) -> JiraIssue:
        async with semaphore:
            return await client.hydrate_issue(issue)

    return list(await asyncio.gather(*(hydrate(issue) for issue in issues)))


def _attach_reverse_children(issues: list[JiraIssue]) -> list[JiraIssue]:
    children_by_parent: dict[str, list[JiraIssueRef]] = {}
    for issue in issues:
        if issue.parent is None:
            continue
        children_by_parent.setdefault(issue.parent.key, []).append(
            JiraIssueRef(
                key=issue.key,
                title=issue.title,
                relationship="child",
                status=issue.status,
            )
        )
    enriched: list[JiraIssue] = []
    for issue in issues:
        children = [*issue.child_issues, *children_by_parent.get(issue.key, [])]
        deduped = {(child.key, child.relationship): child for child in children}
        enriched.append(replace(issue, child_issues=tuple(deduped.values())))
    return enriched


def _parse_issue_ref(value: object, *, relationship: str) -> JiraIssueRef | None:
    row = _dict(value)
    key = _str(row.get("key"))
    if not key:
        return None
    fields = _dict(row.get("fields"))
    status = _dict(fields.get("status"))
    return JiraIssueRef(
        key=key,
        title=_str(fields.get("summary")) or key,
        relationship=relationship,
        status=_str(status.get("name")),
    )


def _parse_issue_links(value: object) -> tuple[JiraIssueRef, ...]:
    refs: list[JiraIssueRef] = []
    for item in _list(value):
        link = _dict(item)
        link_type = _dict(link.get("type"))
        inward = _parse_issue_ref(
            link.get("inwardIssue"),
            relationship=_str(link_type.get("inward")) or "related from",
        )
        outward = _parse_issue_ref(
            link.get("outwardIssue"),
            relationship=_str(link_type.get("outward")) or "relates to",
        )
        if inward is not None:
            refs.append(inward)
        if outward is not None:
            refs.append(outward)
    return tuple(refs)


def _parse_attachment(value: object) -> JiraAttachment | None:
    row = _dict(value)
    attachment_id = _str(row.get("id"))
    filename = _str(row.get("filename"))
    if not attachment_id or not filename:
        return None
    author = _dict(row.get("author"))
    return JiraAttachment(
        id=attachment_id,
        filename=filename,
        mime_type=_str(row.get("mimeType")),
        size_bytes=_int(row.get("size")),
        content_url=_str(row.get("content")),
        created_at=_str(row.get("created")),
        author=_str(author.get("displayName")),
    )


def _parse_comment(row: JsonDict, *, fallback_id: str) -> JiraComment | None:
    body = adf_to_text(row.get("body"))
    if not body:
        return None
    author = _dict(row.get("author"))
    return JiraComment(
        id=_str(row.get("id")) or fallback_id,
        author=_str(author.get("displayName")) or "Unknown author",
        body=body,
        created_at=_str(row.get("created")),
        updated_at=_str(row.get("updated")),
    )


def _attachment_is_extractable(attachment: JiraAttachment) -> bool:
    filename = attachment.filename.lower()
    return filename.endswith(
        (
            ".xlsx",
            ".csv",
            ".docx",
            ".pdf",
            ".png",
            ".jpg",
            ".jpeg",
            ".webp",
            ".txt",
            ".md",
            ".json",
            ".yaml",
            ".yml",
            ".log",
        )
    )


def _extract_attachment_text(filename: str, payload: bytes) -> str:
    suffix = filename.lower()
    if suffix.endswith(".xlsx"):
        return _extract_xlsx_text(payload)
    if suffix.endswith(".docx"):
        return _extract_docx_text(payload)
    if suffix.endswith(".pdf"):
        return _extract_pdf_text(payload)
    if suffix.endswith((".png", ".jpg", ".jpeg", ".webp")):
        return _extract_image_text(payload)
    return payload.decode("utf-8", errors="replace")[:120_000].strip()


def _extract_docx_text(payload: bytes) -> str:
    from docx import Document as WordDocument

    document = WordDocument(io.BytesIO(payload))
    sections: list[str] = []
    sections.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    for table_index, table in enumerate(document.tables, start=1):
        rows = [
            " | ".join(re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells)
            for row in table.rows
        ]
        rows = [row for row in rows if row.strip(" |")]
        if rows:
            sections.extend([f"### Table {table_index}", *rows])
    return "\n\n".join(sections)[:120_000].strip()


def _extract_pdf_text(payload: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(payload))
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)[:120_000].strip()


def _extract_image_text(payload: bytes) -> str:
    import pytesseract
    from PIL import Image

    with Image.open(io.BytesIO(payload)) as image:
        text = cast(str, pytesseract.image_to_string(image))
        return text.strip()[:120_000]


def _extract_xlsx_text(payload: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(payload)) as workbook:
        shared_strings: list[str] = []
        if "xl/sharedStrings.xml" in workbook.namelist():
            root = ElementTree.fromstring(workbook.read("xl/sharedStrings.xml"))
            shared_strings = [
                "".join(node.text or "" for node in item.iter() if node.tag.endswith("}t"))
                for item in root
            ]

        sections: list[str] = []
        sheet_names = sorted(
            name
            for name in workbook.namelist()
            if re.fullmatch(r"xl/worksheets/sheet\d+\.xml", name)
        )
        for sheet_index, sheet_name in enumerate(sheet_names, start=1):
            root = ElementTree.fromstring(workbook.read(sheet_name))
            rows: list[str] = []
            for row in (node for node in root.iter() if node.tag.endswith("}row")):
                cells: list[str] = []
                for cell in (node for node in row if node.tag.endswith("}c")):
                    value_node = next((node for node in cell if node.tag.endswith("}v")), None)
                    inline_nodes = [node for node in cell.iter() if node.tag.endswith("}t")]
                    value = "".join(node.text or "" for node in inline_nodes)
                    if value_node is not None and value_node.text is not None:
                        value = value_node.text
                        if cell.attrib.get("t") == "s" and value.isdigit():
                            index = int(value)
                            if index < len(shared_strings):
                                value = shared_strings[index]
                    cells.append(re.sub(r"\s+", " ", value).strip())
                if any(cells):
                    rows.append(" | ".join(cells))
                if sum(len(item) for item in rows) >= 100_000:
                    break
            if rows:
                sections.append(f"### Worksheet {sheet_index}\n\n" + "\n".join(rows))
        return "\n\n".join(sections)[:120_000].strip()


def _render_issue_markdown(*, board: JiraBoard, issue: JiraIssue) -> str:
    _ = board
    fields = [
        ("Issue key", issue.key),
        ("Project", issue.project_name or issue.project_key),
        ("Type", issue.issue_type),
        ("Status", issue.status),
        ("Priority", issue.priority),
        ("Labels", ", ".join(issue.labels) if issue.labels else None),
        ("Components", ", ".join(issue.components) if issue.components else None),
        ("Assignee", issue.assignee),
        ("Updated", issue.updated_at),
        (
            "Parent epic/issue",
            f"{issue.parent.key}: {issue.parent.title}" if issue.parent is not None else None,
        ),
    ]
    lines = [f"# {issue.key}: {issue.title}", ""]
    lines.extend(f"- **{label}:** {value}" for label, value in fields if value)
    description = _clean_jira_description(issue.description)
    if description:
        lines.extend(["", "## Description", "", description])
    relationships = [*issue.child_issues, *issue.related_issues]
    if relationships:
        lines.extend(["", "## Connected Jira issues", ""])
        lines.extend(
            f"- **{reference.key}: {reference.title}** - {reference.relationship}"
            + (f" - Status: {reference.status}" if reference.status else "")
            for reference in relationships
        )
    if issue.comments:
        lines.extend(["", "## Comments", ""])
        for index, comment in enumerate(issue.comments, start=1):
            timestamp = comment.updated_at or comment.created_at or "undated"
            lines.extend(
                [
                    f"### Comment {index} by {comment.author} ({timestamp})",
                    "",
                    comment.body,
                    "",
                ]
            )
    if issue.attachments:
        lines.extend(["", "## Attachments", ""])
        for attachment in issue.attachments:
            details = [
                attachment.mime_type or "unknown type",
                f"{attachment.size_bytes} bytes" if attachment.size_bytes is not None else "unknown size",
            ]
            lines.extend(
                [
                    f"### Attachment: {attachment.filename}",
                    "",
                    f"- Metadata: {', '.join(details)}",
                    f"- Download: {attachment.content_url}" if attachment.content_url else "",
                ]
            )
            if attachment.extracted_text:
                lines.extend(["", "#### Extracted attachment content", "", attachment.extracted_text])
    return "\n".join(lines).replace("\x00", "").strip() + "\n"


def _issue_metadata(*, board: JiraBoard, issue: JiraIssue) -> dict[str, object]:
    rendered = _render_issue_markdown(board=board, issue=issue).encode("utf-8")
    project = issue.project_key or ""
    owner = issue.assignee_email or issue.assignee or issue.assignee_account_id
    metadata: dict[str, object] = {
        "source": "jira",
        "source_type": "jira",
        "source_family": "jira",
        "source_system": "jira",
        "source_id": issue.key,
        "source_title": f"{issue.key}: {issue.title}",
        "source_url": issue.url,
        "source_space": project,
        "source_version": issue.updated_at,
        "source_updated_at": issue.updated_at,
        "project": project,
        "project_key": issue.project_key,
        "project_name": issue.project_name,
        "issue_id": issue.id,
        "issue_key": issue.key,
        "title": f"{issue.key}: {issue.title}",
        "url": issue.url,
        "created_at": issue.created_at,
        "updated_at": issue.updated_at,
        "status": issue.status,
        "labels": issue.labels,
        "components": issue.components,
        "parent_issue_key": issue.parent.key if issue.parent else None,
        "child_issue_keys": [item.key for item in issue.child_issues],
        "related_issue_keys": [item.key for item in issue.related_issues],
        "comment_count": len(issue.comments),
        "attachment_names": [item.filename for item in issue.attachments],
        "extracted_attachment_count": sum(1 for item in issue.attachments if item.extracted_text),
        "owner": owner,
        "acl": "connector-visible",
        "connector": "jira",
        "connector_scope": project or str(board.id),
        "connector_sync_id": f"jira:{project or 'unknown'}:{issue.key}:{issue.updated_at or issue.id}",
        "chunk_strategy_version": CHUNK_STRATEGY_VERSION,
        "permission_state": "visible",
        "priority": issue.priority,
        "issue_type": issue.issue_type,
        "jira_board_id": board.id,
        "jira_board_name": board.name,
        "jira_project_key": issue.project_key,
        "jira_project_name": issue.project_name,
        "jira_issue_id": issue.id,
        "jira_issue_key": issue.key,
        "jira_issue_url": issue.url,
        "jira_issue_status": issue.status,
        "jira_issue_status_category": issue.status_category,
        "jira_issue_status_category_key": issue.status_category_key,
        "jira_issue_type": issue.issue_type,
        "jira_labels": issue.labels,
        "jira_components": issue.components,
        "jira_parent_issue_key": issue.parent.key if issue.parent else None,
        "jira_child_issue_keys": [item.key for item in issue.child_issues],
        "jira_related_issue_keys": [item.key for item in issue.related_issues],
        "jira_comment_count": len(issue.comments),
        "jira_attachment_names": [item.filename for item in issue.attachments],
        "jira_assignee": issue.assignee,
        "jira_assignee_email": issue.assignee_email,
        "jira_assignee_account_id": issue.assignee_account_id,
        "jira_reporter": issue.reporter,
        "jira_reporter_email": issue.reporter_email,
        "jira_reporter_account_id": issue.reporter_account_id,
        "jira_issue_updated_at": issue.updated_at,
        "source_sha256": hashlib.sha256(rendered).hexdigest(),
    }
    return normalize_source_metadata(
        metadata,
        source_type="jira",
        title=f"{issue.key}: {issue.title}",
        source_id=issue.key,
        source_url=issue.url,
        source_space=project,
        source_version=issue.updated_at,
        updated_at=issue.updated_at,
        status=issue.status,
        labels=issue.labels,
        owner=owner,
        acl="connector-visible",
        connector="jira",
        connector_scope=project or str(board.id),
        source_sha256=str(metadata["source_sha256"]),
    )


def _jira_issue_allowed(issue: JiraIssue, settings: Settings) -> bool:
    status = (issue.status or "").lower()
    labels = {label.lower() for label in issue.labels}
    issue_type = (issue.issue_type or "").lower()
    include_statuses = _csv_set(settings.jira_include_statuses)
    exclude_statuses = _csv_set(settings.jira_exclude_statuses)
    include_labels = _csv_set(settings.jira_include_labels)
    exclude_labels = _csv_set(settings.jira_exclude_labels)
    include_issue_types = _csv_set(settings.jira_include_issue_types)
    exclude_issue_types = _csv_set(settings.jira_exclude_issue_types)
    if include_statuses and status not in include_statuses:
        return False
    if exclude_statuses and status in exclude_statuses:
        return False
    if include_labels and not (labels & include_labels):
        return False
    if exclude_labels and labels & exclude_labels:
        return False
    if include_issue_types and issue_type not in include_issue_types:
        return False
    if exclude_issue_types and issue_type in exclude_issue_types:
        return False
    return not (settings.jira_drop_empty_descriptions and not _clean_jira_description(issue.description))


def _csv_set(value: str) -> set[str]:
    return {item.strip().lower() for item in value.split(",") if item.strip()}


def _issue_filename(issue: JiraIssue) -> str:
    stem = re.sub(r"[^A-Za-z0-9._ -]+", "-", f"{issue.key}-{issue.title}").strip(" .-")
    stem = re.sub(r"\s+", "-", stem)[:100] or issue.key
    return f"{stem}.md"


def _is_current(existing: Document | None, metadata: dict[str, object]) -> bool:
    if existing is None or existing.status != DocumentStatus.READY:
        return False
    existing_metadata = existing.doc_metadata or {}
    return (
        existing_metadata.get("source_sha256") == metadata.get("source_sha256")
        and existing_metadata.get("chunk_strategy_version") == CHUNK_STRATEGY_VERSION
    )


def _dict(value: object) -> JsonDict:
    return cast(JsonDict, value) if isinstance(value, dict) else {}


def _list(value: object) -> list[object]:
    return value if isinstance(value, list) else []


def _list_of_dicts(value: object) -> list[JsonDict]:
    if not isinstance(value, list):
        return []
    return [cast(JsonDict, item) for item in value if isinstance(item, dict)]


def _required_str(row: JsonDict, key: str) -> str:
    value = _str(row.get(key))
    if not value:
        raise ProviderError(f"Jira response is missing {key}")
    return value


def _str(value: object) -> str | None:
    return value if isinstance(value, str) else None


def _int(value: object) -> int | None:
    return value if isinstance(value, int) else None
